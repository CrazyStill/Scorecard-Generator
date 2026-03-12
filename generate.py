import os, csv, shutil, tempfile
from docx import Document
import docx2pdf
from PyPDF2 import PdfMerger
import comtypes.client
import pythoncom


# ── Paragraph / run helpers ───────────────────────────────────────────────────

def merge_runs_in_paragraph(paragraph):
    """Collapse all runs in a paragraph into the first run.

    python-docx can split a single word across multiple runs when the DOCX
    has mixed formatting (e.g. spell-check marks). Merging before replacement
    ensures placeholder strings are never split across run boundaries.
    """
    if len(paragraph.runs) <= 1:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = full_text


def replace_text_in_paragraph(paragraph, placeholder, replacement):
    """Replace a single placeholder token in a paragraph's first run.

    Safe to call after merge_runs_in_paragraph — the merged text lives in
    runs[0], so a direct string replace is sufficient.
    """
    if placeholder in paragraph.text:
        new_text = paragraph.text.replace(placeholder, replacement)
        paragraph.runs[0].text = new_text


def replace_placeholders_in_paragraphs(paragraphs, placeholders):
    """Apply all placeholder replacements across a list of paragraphs."""
    for paragraph in paragraphs:
        merge_runs_in_paragraph(paragraph)
        for placeholder, replacement in placeholders:
            replace_text_in_paragraph(paragraph, placeholder, replacement)


# ── Document-level replacement ────────────────────────────────────────────────

def replace_text_in_doc(doc, rows, mapping, cards_per_page=4):
    """Fill all placeholder tokens in a DOCX document with CSV row values.

    Placeholders follow the convention  <base_name>_<card_number>  (e.g.
    NAME_1, NAME_2). Each card slot maps to one row in `rows`; slots without
    a corresponding row get an empty string so the template stays clean.

    Args:
        doc:           python-docx Document object (modified in place).
        rows:          List of dicts from csv.DictReader for this page group.
        mapping:       Dict of {csv_header: placeholder_base_name}.
        cards_per_page: Number of card slots on the template page.
    """
    all_placeholders = []
    for i in range(1, cards_per_page + 1):
        row = rows[i - 1] if i - 1 < len(rows) else None
        for csv_header, placeholder in mapping.items():
            value = row.get(csv_header, "") if row else ""
            all_placeholders.append((f"{placeholder}_{i}", value))

    # Replace in body paragraphs and in every table cell
    replace_placeholders_in_paragraphs(doc.paragraphs, all_placeholders)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholders_in_paragraphs(cell.paragraphs, all_placeholders)


# ── PDF utilities ─────────────────────────────────────────────────────────────

def merge_two_pdfs(pdf1, pdf2, merged_pdf):
    """Concatenate pdf1 and pdf2 into merged_pdf using PyPDF2."""
    merger = PdfMerger()
    merger.append(pdf1)
    merger.append(pdf2)
    merger.write(merged_pdf)
    merger.close()


def convert_docx_to_pdf(docx_path, pdf_path):
    """Convert a DOCX file to PDF using Microsoft Word via COM automation.

    Tries docx2pdf first (lighter wrapper). Falls back to direct COM
    automation via comtypes if docx2pdf fails (e.g. in packaged .exe builds).
    COM must be initialised per-thread, hence the CoInitialize / CoUninitialize
    calls — this function is safe to call from background threads.
    """
    pythoncom.CoInitialize()
    try:
        docx2pdf.convert(docx_path, pdf_path)
    except Exception:
        # Direct COM fallback: open Word invisibly, save as PDF, then quit
        wdFormatPDF = 17
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        try:
            doc.SaveAs2(pdf_path, FileFormat=wdFormatPDF)
        finally:
            doc.Close()
            word.Quit()
    finally:
        pythoncom.CoUninitialize()


# ── Main generation pipeline ──────────────────────────────────────────────────

def generate_scorecard(template_path, csv_path, mapping,
                       cards_per_page=4, back_pdf_path=None, temp_dir=None,
                       progress_callback=None):
    """Generate a merged scorecard PDF from a DOCX template and a CSV data file.

    For each group of `cards_per_page` CSV rows:
      1. Load the DOCX template and fill in the placeholder tokens.
      2. Convert the filled DOCX to a PDF page via Word COM automation.
      3. Optionally append a back-design PDF page.

    All intermediate files are written to `temp_dir`. The final merged PDF is
    returned as a path inside `temp_dir` — the caller is responsible for
    moving or cleaning it up.

    Args:
        template_path:     Path to the front DOCX template.
        csv_path:          Path to the filled CSV data file.
        mapping:           Dict of {csv_header: placeholder_base_name}.
        cards_per_page:    Number of scorecards per printed page (1–4).
        back_pdf_path:     Optional path to a back-design PDF page.
        temp_dir:          Working directory for intermediate files.
        progress_callback: Optional callable(step, total, message) called
                           after each page is converted so callers can show
                           a progress bar.

    Returns:
        Absolute path to the final merged PDF inside temp_dir.

    Raises:
        ValueError: If the CSV contains no data rows.
    """
    final_pdf_list = []

    # Read all data rows up front so we know the total page count before
    # the loop starts (needed for accurate progress reporting).
    with open(csv_path, newline='', encoding='latin-1') as f:
        sample = f.read(1024)
        f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample)
        except csv.Error:
            # Sniffer fails on very short or simple files; fall back to Excel CSV
            dialect = csv.excel
        reader = csv.DictReader(f, dialect=dialect)
        # Skip entirely blank rows so empty trailing lines don't create blank pages
        rows = [row for row in reader if any(value.strip() for value in row.values())]

    if not rows:
        raise ValueError("The CSV file contains no data rows. Please fill in the template and try again.")

    total_pages = len(range(0, len(rows), cards_per_page))

    for i in range(0, len(rows), cards_per_page):
        group = rows[i : i + cards_per_page]
        idx   = i // cards_per_page

        # Fill placeholders in a fresh copy of the template for each page
        front_doc = Document(template_path)
        replace_text_in_doc(front_doc, group, mapping, cards_per_page=cards_per_page)

        temp_front_docx = os.path.join(temp_dir, f"temp_front_{idx}.docx")
        temp_front_pdf  = os.path.join(temp_dir, f"temp_front_{idx}.pdf")
        page_pdf        = os.path.join(temp_dir, f"page_{idx}.pdf")

        front_doc.save(temp_front_docx)
        convert_docx_to_pdf(temp_front_docx, temp_front_pdf)  # slow — Word COM call

        if back_pdf_path and os.path.exists(back_pdf_path):
            # Merge front + back into a single two-page sheet for this card group
            merge_two_pdfs(temp_front_pdf, back_pdf_path, page_pdf)
        else:
            shutil.copy(temp_front_pdf, page_pdf)

        final_pdf_list.append(page_pdf)

        # Clean up per-page intermediates immediately to keep temp_dir small
        os.remove(temp_front_docx)
        os.remove(temp_front_pdf)

        if progress_callback:
            progress_callback(idx + 1, total_pages, f"Converting page {idx + 1} of {total_pages}…")

    # Merge all page PDFs into one final document in page-number order
    output_pdf = os.path.join(temp_dir, "Final_Scorecards.pdf")
    merger = PdfMerger()
    for pdf in sorted(final_pdf_list, key=lambda x: int(x.split('_')[-1].split('.')[0])):
        merger.append(pdf)
    merger.write(output_pdf)
    merger.close()

    # Remove individual page PDFs now that they're merged
    for pdf in final_pdf_list:
        os.remove(pdf)

    return output_pdf
